# import streamlit as st
# import pdfplumber
# import docx
# import re
# import os
# import logging

# # Setup basic logging to help debug in the terminal
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# # ==============================================================================
# # PARSER 1: FOR PDFs WITH "Ans." line
# # ==============================================================================
# def parse_pdf_with_ans_lines(pdf_path):
#     logging.info("Using Parser 1 for PDFs with 'Ans.' lines.")
#     full_text = ""
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text(x_tolerance=1)
#                 if page_text:
#                     full_text += page_text + "\n"
#     except Exception as e:
#         st.error(f"Error reading PDF file: {e}")
#         logging.error(f"Failed to read PDF: {e}")
#         return []

#     full_text = re.sub(r'Question Paper – \d+', '', full_text)
#     full_text = re.sub(r'^\s*\d+\s*$', '', full_text, flags=re.MULTILINE)
#     questions_data = []
    
#     question_blocks = re.split(r'\n(?=\d{1,3}\.\s)', full_text)
    
#     for block in question_blocks:
#         block = block.strip()
#         if not block: continue

#         answer_match = re.search(r'Ans\.\s*\(?([a-zA-Z])\)?', block)
#         if not answer_match: continue
        
#         correct_letter = answer_match.group(1).lower()
#         content = block[:answer_match.start()].strip()
        
#         lines = content.split('\n')
#         question_text, options_text = "", ""
#         options_started = False

#         for line in lines:
#             if re.match(r'^\s*[a-d][.)]', line.strip()):
#                 options_started = True
#             if options_started:
#                 options_text += " " + line
#             else:
#                 question_text += " " + line
        
#         question_text = re.sub(r'^\d{1,3}\.\s*', '', question_text).strip()
#         options = re.findall(r'([a-d])[.)]\s*(.*?)(?=\s*[a-d][.)]|$)', options_text, re.DOTALL)
#         opts_dict = {key.lower(): " ".join(val.strip().split()) for key, val in options}

#         if correct_letter in opts_dict and question_text:
#             questions_data.append({
#                 "question": question_text,
#                 "options": [opts_dict.get(c, '') for c in 'abcd'],
#                 "answer": opts_dict[correct_letter]
#             })
            
#     logging.info(f"Parser 1 found {len(questions_data)} questions.")
#     return questions_data

# # ==============================================================================
# # PARSER 2: FOR PDFs WITH AN ANSWER TABLE AT THE END
# # ==============================================================================
# def parse_pdf_with_answer_table(pdf_path):
#     logging.info("Using Parser 2 for PDFs with an answer table.")
#     questions, answer_map = {}, {}
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             full_text = "".join(page.extract_text() + "\n" for page in pdf.pages if page.extract_text())
        
#         answer_section_start = re.search(r'\nANSWERS\s*\n', full_text, re.IGNORECASE)
#         main_content = full_text[:answer_section_start.start()] if answer_section_start else ""
#         answer_content = full_text[answer_section_start.start():] if answer_section_start else ""

#         if not main_content or not answer_content:
#             logging.error("Could not separate main content from the answer section.")
#             return []

#         question_blocks = re.split(r'\n(?=\d{1,3}\.\s)', main_content)

#         for block in question_blocks:
#             block = block.strip()
#             q_match = re.match(r'^(\d{1,3})\.\s+(.*)', block, re.DOTALL)
#             if not q_match: continue

#             q_num, content = q_match.groups()
#             lines = content.split('\n')
#             question_text, options_list, is_question_part = "", [], True
            
#             for line in lines:
#                 line = line.strip()
#                 if re.match(r'^[a-c]\)\s', line):
#                     is_question_part = False
#                     options_list.append(re.sub(r'^[a-c]\)\s*', '', line))
#                 elif is_question_part:
#                     question_text += " " + line
#                 elif options_list:
#                     options_list[-1] += " " + line
            
#             if question_text and len(options_list) > 0:
#                 questions[q_num] = { "question": " ".join(question_text.strip().split()), "options": [opt.strip() for opt in options_list] }
        
#         answer_pairs = re.findall(r'(\d+)\s+([A-C])', answer_content)
#         answer_map = {num: letter for num, letter in answer_pairs}

#         if not answer_map:
#             logging.error("Could not parse the answer table.")
#             return []

#         final_data = []
#         for q_num, q_data in questions.items():
#             ans_letter = answer_map.get(q_num)
#             if ans_letter:
#                 ans_index = ord(ans_letter.lower()) - 97
#                 if ans_index < len(q_data['options']):
#                     q_data['answer'] = q_data['options'][ans_index]
#                     final_data.append(q_data)
        
#         logging.info(f"Parser 2 found {len(final_data)} questions.")
#         return final_data
#     except Exception as e:
#         st.error(f"An error occurred during PDF parsing: {e}")
#         logging.error(f"Parser 2 failed: {e}")
#         return []

# # ==============================================================================
# # PARSER 3: ROBUST DOCX PARSER WITH BOLD DETECTION (Completely Rewritten)
# # ==============================================================================
# def parse_docx_with_bold_answers(docx_path):
#     logging.info("LOG: Using NEW robust Parser 3 for DOCX.")
#     try:
#         doc = docx.Document(docx_path)
#         questions_data = []
#         current_q_block_paras = []

#         for para in doc.paragraphs:
#             para_text = para.text.strip()
#             # Skip empty paragraphs and headers
#             if not para_text or para_text == "Technical general- test 02":
#                 continue

#             # Check if a paragraph is a new question by its style
#             # This is the key insight from the debug file
#             if para.style.name == 'List Paragraph':
#                 # If we have a block of paragraphs collected, process them
#                 if current_q_block_paras:
#                     question = process_docx_block(current_q_block_paras)
#                     if question:
#                         questions_data.append(question)
#                 # Start a new block with the current paragraph
#                 current_q_block_paras = [para]
#             elif current_q_block_paras:
#                 # If it's not a new question, it's a continuation of the previous one
#                 current_q_block_paras.append(para)

#         # Process the very last block in the file
#         if current_q_block_paras:
#             question = process_docx_block(current_q_block_paras)
#             if question:
#                 questions_data.append(question)
        
#         logging.info(f"Parser 3 found {len(questions_data)} questions.")
#         return questions_data

#     except Exception as e:
#         st.error(f"Error reading DOCX file: {e}")
#         logging.error(f"Parser 3 failed: {e}")
#         return []

# def process_docx_block(paragraphs):
#     """
#     Helper function to process a single block of DOCX paragraphs into one question.
#     This now correctly identifies the question and options based on position.
#     """
#     if not paragraphs:
#         return None

#     # The first paragraph is ALWAYS the question text
#     question_text = paragraphs[0].text.strip()
    
#     options = []
#     correct_answer = ""

#     # All subsequent paragraphs are options
#     for para in paragraphs[1:]:
#         option_text = para.text.strip()
#         if not option_text: continue
        
#         options.append(option_text)
        
#         # Check if any run in this paragraph is bold
#         if any(run.bold for run in para.runs):
#             correct_answer = option_text
            
#     if question_text and len(options) > 0 and correct_answer:
#         return {
#             "question": question_text,
#             "options": options,
#             "answer": correct_answer
#         }
#     return None

# # ==============================================================================
# # STREAMLIT APPLICATION UI
# # ==============================================================================
# def main():
#     st.set_page_config(page_title="Quiz Generator", page_icon="✈️", layout="centered")
#     st.title("✈️ Universal Quiz Generator")

#     if 'state' not in st.session_state:
#         st.session_state.state = 'initial'
#         st.session_state.questions = []
#         st.session_state.current_question = 0
#         st.session_state.user_answers = []

#     if st.session_state.state == 'initial':
#         st.info("Select your quiz format, then upload the corresponding file.")
        
#         parser_choice = st.radio(
#             "**1. Choose your file's format:**",
#             ('**Pattern 1 (PDF):** Answer is `Ans. a` after each question.', 
#              '**Pattern 2 (PDF):** Answers are in a table at the end.',
#              '**Pattern 3 (DOCX):** The correct answer is **bolded**.')
#         )

#         file_type = "docx" if "Pattern 3" in parser_choice else "pdf"
#         uploaded_file = st.file_uploader(f"**2. Upload your .{file_type} file:**", type=[file_type])

#         if uploaded_file:
#             temp_dir = "uploads"
#             os.makedirs(temp_dir, exist_ok=True)
#             temp_file_path = os.path.join(temp_dir, uploaded_file.name)
#             with open(temp_file_path, "wb") as f:
#                 f.write(uploaded_file.getbuffer())

#             if 'Pattern 1' in parser_choice:
#                 st.session_state.questions = parse_pdf_with_ans_lines(temp_file_path)
#             elif 'Pattern 2' in parser_choice:
#                 st.session_state.questions = parse_pdf_with_answer_table(temp_file_path)
#             else: # Pattern 3
#                 st.session_state.questions = parse_docx_with_bold_answers(temp_file_path)
            
#             if st.session_state.questions:
#                 st.session_state.state = 'quiz_started'
#                 st.rerun()
#             else:
#                 st.error("Could not find any valid questions using the selected parser. Please check the file and your format selection.")
#         return

#     if st.session_state.state == 'finished':
#         score = sum(1 for i, q in enumerate(st.session_state.questions) if i < len(st.session_state.user_answers) and st.session_state.user_answers[i] == q['answer'])
        
#         st.success(f"## 🎯 Quiz Complete!")
#         st.write(f"### Your Final Score: {score} out of {len(st.session_state.questions)}")
#         with st.expander("Review Your Answers"):
#              for i, q in enumerate(st.session_state.questions):
#                 user_answer = st.session_state.user_answers[i]
#                 correct_answer = q['answer']
#                 if user_answer == correct_answer:
#                     st.markdown(f"**Q{i+1}: {q['question']}**\n\n✅ Your answer: `{user_answer}` (Correct)")
#                 else:
#                     st.markdown(f"**Q{i+1}: {q['question']}**\n\n❌ Your answer: `{user_answer}`\n\nCorrect answer: `{correct_answer}`")
#                 st.markdown("---")
#         if st.button("Take a New Quiz"):
#             for key in list(st.session_state.keys()): del st.session_state[key]
#             st.rerun()
#         return

#     if st.session_state.state in ['quiz_started', 'show_feedback']:
#         q_index = st.session_state.current_question
#         q_data = st.session_state.questions[q_index]

#         st.subheader(f"Question {q_index + 1} of {len(st.session_state.questions)}")
#         st.markdown(f"<p style='font-size: 20px; font-weight: 500;'>{q_data['question']}</p>", unsafe_allow_html=True)
        
#         valid_options = [opt for opt in q_data["options"] if opt]
        
#         if st.session_state.state == 'quiz_started':
#             with st.form(key=f"form_{q_index}"):
#                 user_choice = st.radio("Choose your answer:", options=valid_options, key=f"radio_{q_index}")
#                 submitted = st.form_submit_button("Check Answer")
#                 if submitted:
#                     st.session_state.user_answers.append(user_choice)
#                     st.session_state.state = 'show_feedback'
#                     st.rerun()
        
#         elif st.session_state.state == 'show_feedback':
#             last_answer = st.session_state.user_answers[q_index]
#             try: default_index = valid_options.index(last_answer)
#             except (ValueError, IndexError): default_index = 0
            
#             st.radio("Your Answer:", options=valid_options, index=default_index, disabled=True)
            
#             correct_answer = q_data['answer']
#             if last_answer == correct_answer:
#                 st.success("✅ Correct!")
#             else:
#                 st.error(f"❌ Incorrect! The correct answer was: **{correct_answer}**")
            
#             button_text = "Finish Quiz" if q_index + 1 == len(st.session_state.questions) else "Next Question"
#             if st.button(button_text):
#                 if q_index + 1 < len(st.session_state.questions):
#                     st.session_state.current_question += 1
#                     st.session_state.state = 'quiz_started'
#                 else:
#                     st.session_state.state = 'finished'
#                 st.rerun()

# if __name__ == "__main__":
#     main()







################################################################
# import streamlit as st
# import pdfplumber
# import docx
# import re
# import os
# import logging

# # Setup basic logging to help debug in the terminal
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# # ==============================================================================
# # PARSER 1: FOR PDFs WITH "Ans." line (Working and Stable)
# # ==============================================================================
# def parse_pdf_with_ans_lines(pdf_path):
#     logging.info("Using Parser 1 for PDFs with 'Ans.' lines.")
#     full_text = ""
#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             for page in pdf.pages:
#                 page_text = page.extract_text(x_tolerance=1)
#                 if page_text:
#                     full_text += page_text + "\n"
#     except Exception as e:
#         st.error(f"Error reading PDF file: {e}")
#         logging.error(f"Failed to read PDF: {e}")
#         return []

#     # Clean the text
#     full_text = re.sub(r'Question Paper – \d+', '', full_text)
#     full_text = re.sub(r'^\s*\d+\s*$', '', full_text, flags=re.MULTILINE)
#     questions_data = []
    
#     question_blocks = re.split(r'\n(?=\d{1,3}\.\s)', full_text)
    
#     for block in question_blocks:
#         block = block.strip()
#         if not block: continue

#         answer_match = re.search(r'Ans\.\s*\(?([a-zA-Z])\)?', block)
#         if not answer_match: continue
        
#         correct_letter = answer_match.group(1).lower()
#         content = block[:answer_match.start()].strip()
        
#         lines = content.split('\n')
#         question_text, options_text = "", ""
#         options_started = False

#         for line in lines:
#             if re.match(r'^\s*[a-d][.)]', line.strip()):
#                 options_started = True
#             if options_started:
#                 options_text += " " + line
#             else:
#                 question_text += " " + line
        
#         question_text = re.sub(r'^\d{1,3}\.\s*', '', question_text).strip()
#         options = re.findall(r'([a-d])[.)]\s*(.*?)(?=\s*[a-d][.)]|$)', options_text, re.DOTALL)
#         opts_dict = {key.lower(): " ".join(val.strip().split()) for key, val in options}

#         if correct_letter in opts_dict and question_text:
#             questions_data.append({
#                 "question": question_text,
#                 "options": [opts_dict.get(c, '') for c in 'abcd'],
#                 "answer": opts_dict[correct_letter]
#             })
            
#     logging.info(f"Parser 1 found {len(questions_data)} questions.")
#     return questions_data

# # ==============================================================================
# # PARSER 2: FOR PDFs WITH AN ANSWER TABLE AT THE END (Working and Stable)
# # ==============================================================================
# def parse_pdf_with_answer_table(pdf_path):
#     logging.info("Using Parser 2 for PDFs with an answer table.")
#     questions = {}
#     answer_map = {}

#     try:
#         with pdfplumber.open(pdf_path) as pdf:
#             full_text = "".join(page.extract_text() + "\n" for page in pdf.pages if page.extract_text())
        
#         answer_section_start = re.search(r'\nANSWERS\s*\n', full_text, re.IGNORECASE)
#         main_content = full_text[:answer_section_start.start()] if answer_section_start else ""
#         answer_content = full_text[answer_section_start.start():] if answer_section_start else ""

#         if not main_content or not answer_content:
#             logging.error("Could not separate main content from the answer section.")
#             return []

#         question_blocks = re.split(r'\n(?=\d{1,3}\.\s)', main_content)

#         for block in question_blocks:
#             block = block.strip()
#             q_match = re.match(r'^(\d{1,3})\.\s+(.*)', block, re.DOTALL)
#             if not q_match: continue

#             q_num, content = q_match.groups()
#             lines = content.split('\n')
#             question_text, options_list, is_question_part = "", [], True
            
#             for line in lines:
#                 line = line.strip()
#                 if re.match(r'^[a-c]\)\s', line):
#                     is_question_part = False
#                     options_list.append(re.sub(r'^[a-c]\)\s*', '', line))
#                 elif is_question_part:
#                     question_text += " " + line
#                 elif options_list:
#                     options_list[-1] += " " + line
            
#             if question_text and len(options_list) > 0:
#                 questions[q_num] = { "question": " ".join(question_text.strip().split()), "options": [opt.strip() for opt in options_list] }
        
#         answer_pairs = re.findall(r'(\d+)\s+([A-C])', answer_content)
#         answer_map = {num: letter for num, letter in answer_pairs}

#         if not answer_map:
#             logging.error("Could not parse the answer table.")
#             return []

#         final_data = []
#         for q_num, q_data in questions.items():
#             ans_letter = answer_map.get(q_num)
#             if ans_letter:
#                 ans_index = ord(ans_letter.lower()) - 97
#                 if ans_index < len(q_data['options']):
#                     q_data['answer'] = q_data['options'][ans_index]
#                     final_data.append(q_data)
        
#         logging.info(f"Parser 2 found {len(final_data)} questions.")
#         return final_data

#     except Exception as e:
#         st.error(f"An error occurred during PDF parsing: {e}")
#         logging.error(f"Parser 2 failed: {e}")
#         return []

# # ==============================================================================
# # PARSER 3: ROBUST DOCX PARSER WITH BOLD DETECTION (Completely Rewritten)
# # ==============================================================================
# def parse_docx_with_bold_answers(docx_path):
#     logging.info("LOG: Using NEW robust Parser 3 for DOCX.")
#     try:
#         doc = docx.Document(docx_path)
#         questions_data = []
#         current_q_block = []

#         for para in doc.paragraphs:
#             para_text = para.text.strip()
#             if not para_text or "technical test" in para_text.lower():
#                 continue

#             # A new question always starts with a number followed by a dot.
#             if re.match(r'^\d{1,3}\.\s', para_text):
#                 if current_q_block:
#                     question = process_docx_block(current_q_block)
#                     if question:
#                         questions_data.append(question)
#                 current_q_block = [para]
#             elif current_q_block:
#                 current_q_block.append(para)

#         if current_q_block: # Process the last block
#             question = process_docx_block(current_q_block)
#             if question:
#                 questions_data.append(question)
        
#         logging.info(f"Parser 3 found {len(questions_data)} questions.")
#         return questions_data

#     except Exception as e:
#         st.error(f"Error reading DOCX file: {e}")
#         logging.error(f"Parser 3 failed: {e}")
#         return []

# def process_docx_block(paragraphs):
#     """Helper function to process a single block of DOCX paragraphs."""
#     if not paragraphs: return None

#     # Join the text of all paragraphs in the block to handle multi-line questions/options
#     full_block_text = "\n".join([p.text.strip() for p in paragraphs])
    
#     # Extract question text (everything before the first option)
#     first_option_match = re.search(r'\n[a-zA-Z][.)]', full_block_text)
#     if not first_option_match: return None
    
#     question_text = full_block_text[:first_option_match.start()].strip()
#     question_text = re.sub(r'^\d{1,3}\.\s*', '', question_text).replace('\n', ' ')

#     options_text = full_block_text[first_option_match.start():].strip()
    
#     # Extract options using regex
#     options_list = [opt.strip() for opt in re.split(r'[a-zA-Z][.)]\s*', options_text) if opt.strip()]
    
#     # Find the correct answer by checking for bold formatting in the original paragraphs
#     correct_answer = ""
#     for para in paragraphs:
#         is_bold = any(run.bold for run in para.runs)
#         if is_bold:
#             # The text of the bold paragraph is the answer
#             correct_answer = re.sub(r'^[a-zA-Z][.)]\s*', '', para.text).strip()
#             break
            
#     if question_text and len(options_list) > 0 and correct_answer:
#         return {
#             "question": " ".join(question_text.split()),
#             "options": options_list,
#             "answer": correct_answer
#         }
    
#     logging.warning(f"Skipping invalid DOCX block. Question: '{question_text[:50]}...'")
#     return None

# # ==============================================================================
# # STREAMLIT APPLICATION UI
# # ==============================================================================
# def main():
#     st.set_page_config(page_title="Quiz Generator", page_icon="✈️", layout="centered")
#     st.title("✈️ Universal Quiz Generator")

#     if 'state' not in st.session_state:
#         st.session_state.state = 'initial'
#         st.session_state.questions = []
#         st.session_state.current_question = 0
#         st.session_state.user_answers = []

#     if st.session_state.state == 'initial':
#         st.info("Select your quiz format, then upload the corresponding file.")
        
#         parser_choice = st.radio(
#             "**1. Choose your file's format:**",
#             ('**Pattern 1 (PDF):** Answer is `Ans. a` after each question.', 
#              '**Pattern 2 (PDF):** Answers are in a table at the end.',
#              '**Pattern 3 (DOCX):** The correct answer is **bolded**.')
#         )

#         file_type = "docx" if "Pattern 3" in parser_choice else "pdf"
#         uploaded_file = st.file_uploader(f"**2. Upload your .{file_type} file:**", type=[file_type])

#         if uploaded_file:
#             temp_dir = "uploads"
#             os.makedirs(temp_dir, exist_ok=True)
#             temp_file_path = os.path.join(temp_dir, uploaded_file.name)
#             with open(temp_file_path, "wb") as f:
#                 f.write(uploaded_file.getbuffer())

#             if 'Pattern 1' in parser_choice:
#                 st.session_state.questions = parse_pdf_with_ans_lines(temp_file_path)
#             elif 'Pattern 2' in parser_choice:
#                 st.session_state.questions = parse_pdf_with_answer_table(temp_file_path)
#             else: # Pattern 3
#                 st.session_state.questions = parse_docx_with_bold_answers(temp_file_path)
            
#             if st.session_state.questions:
#                 st.session_state.state = 'quiz_started'
#                 st.rerun()
#             else:
#                 st.error("Could not find any valid questions using the selected parser. Please check the file and your format selection.")
#         return

#     if st.session_state.state == 'finished':
#         score = sum(1 for i, q in enumerate(st.session_state.questions) if i < len(st.session_state.user_answers) and st.session_state.user_answers[i] == q['answer'])
        
#         st.success(f"## 🎯 Quiz Complete!")
#         st.write(f"### Your Final Score: {score} out of {len(st.session_state.questions)}")
#         with st.expander("Review Your Answers"):
#              for i, q in enumerate(st.session_state.questions):
#                 user_answer = st.session_state.user_answers[i]
#                 correct_answer = q['answer']
#                 if user_answer == correct_answer:
#                     st.markdown(f"**Q{i+1}: {q['question']}**\n\n✅ Your answer: `{user_answer}` (Correct)")
#                 else:
#                     st.markdown(f"**Q{i+1}: {q['question']}**\n\n❌ Your answer: `{user_answer}`\n\nCorrect answer: `{correct_answer}`")
#                 st.markdown("---")
#         if st.button("Take a New Quiz"):
#             for key in list(st.session_state.keys()): del st.session_state[key]
#             st.rerun()
#         return

#     if st.session_state.state in ['quiz_started', 'show_feedback']:
#         q_index = st.session_state.current_question
#         q_data = st.session_state.questions[q_index]

#         st.subheader(f"Question {q_index + 1} of {len(st.session_state.questions)}")
#         st.markdown(f"<p style='font-size: 20px; font-weight: 500;'>{q_data['question']}</p>", unsafe_allow_html=True)
        
#         valid_options = [opt for opt in q_data["options"] if opt]
        
#         if st.session_state.state == 'quiz_started':
#             with st.form(key=f"form_{q_index}"):
#                 user_choice = st.radio("Choose your answer:", options=valid_options, key=f"radio_{q_index}")
#                 submitted = st.form_submit_button("Check Answer")
#                 if submitted:
#                     st.session_state.user_answers.append(user_choice)
#                     st.session_state.state = 'show_feedback'
#                     st.rerun()
        
#         elif st.session_state.state == 'show_feedback':
#             last_answer = st.session_state.user_answers[q_index]
#             try: default_index = valid_options.index(last_answer)
#             except (ValueError, IndexError): default_index = 0
            
#             st.radio("Your Answer:", options=valid_options, index=default_index, disabled=True)
            
#             correct_answer = q_data['answer']
#             if last_answer == correct_answer:
#                 st.success("✅ Correct!")
#             else:
#                 st.error(f"❌ Incorrect! The correct answer was: **{correct_answer}**")
            
#             button_text = "Finish Quiz" if q_index + 1 == len(st.session_state.questions) else "Next Question"
#             if st.button(button_text):
#                 if q_index + 1 < len(st.session_state.questions):
#                     st.session_state.current_question += 1
#                     st.session_state.state = 'quiz_started'
#                 else:
#                     st.session_state.state = 'finished'
#                 st.rerun()

# if __name__ == "__main__":
#     main()

##########################################
import streamlit as st
import pdfplumber
import docx
import re
import os
import logging

# Setup basic logging to help debug in the terminal
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# ==============================================================================
# PARSER 1: FOR PDFs WITH "Ans." line (Working and Stable)
# ==============================================================================
def parse_pdf_with_ans_lines(pdf_path):
    logging.info("Using Parser 1 for PDFs with 'Ans.' lines.")
    full_text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=1)
                if page_text:
                    full_text += page_text + "\n"
    except Exception as e:
        st.error(f"Error reading PDF file: {e}")
        logging.error(f"Failed to read PDF: {e}")
        return []

    # Clean the text
    full_text = re.sub(r'Question Paper – \d+', '', full_text)
    full_text = re.sub(r'^\s*\d+\s*$', '', full_text, flags=re.MULTILINE)
    questions_data = []
    
    question_blocks = re.split(r'\n(?=\d{1,3}\.\s)', full_text)
    
    for block in question_blocks:
        block = block.strip()
        if not block: continue

        answer_match = re.search(r'Ans\.\s*\(?([a-zA-Z])\)?', block)
        if not answer_match: continue
        
        correct_letter = answer_match.group(1).lower()
        content = block[:answer_match.start()].strip()
        
        lines = content.split('\n')
        question_text, options_text = "", ""
        options_started = False

        for line in lines:
            if re.match(r'^\s*[a-d][.)]', line.strip()):
                options_started = True
            if options_started:
                options_text += " " + line
            else:
                question_text += " " + line
        
        question_text = re.sub(r'^\d{1,3}\.\s*', '', question_text).strip()
        options = re.findall(r'([a-d])[.)]\s*(.*?)(?=\s*[a-d][.)]|$)', options_text, re.DOTALL)
        opts_dict = {key.lower(): " ".join(val.strip().split()) for key, val in options}

        if correct_letter in opts_dict and question_text:
            questions_data.append({
                "question": question_text,
                "options": [opts_dict.get(c, '') for c in 'abcd'],
                "answer": opts_dict[correct_letter]
            })
            
    logging.info(f"Parser 1 found {len(questions_data)} questions.")
    return questions_data

# ==============================================================================
# PARSER 2: FOR PDFs WITH AN ANSWER TABLE AT THE END (Working and Stable)
# ==============================================================================
def parse_pdf_with_answer_table(pdf_path):
    logging.info("Using Parser 2 for PDFs with an answer table.")
    questions = {}
    answer_map = {}

    try:
        with pdfplumber.open(pdf_path) as pdf:
            full_text = "".join(page.extract_text() + "\n" for page in pdf.pages if page.extract_text())
        
        answer_section_start = re.search(r'\nANSWERS\s*\n', full_text, re.IGNORECASE)
        main_content = full_text[:answer_section_start.start()] if answer_section_start else ""
        answer_content = full_text[answer_section_start.start():] if answer_section_start else ""

        if not main_content or not answer_content:
            logging.error("Could not separate main content from the answer section.")
            return []

        question_blocks = re.split(r'\n(?=\d{1,3}\.\s)', main_content)

        for block in question_blocks:
            block = block.strip()
            q_match = re.match(r'^(\d{1,3})\.\s+(.*)', block, re.DOTALL)
            if not q_match: continue

            q_num, content = q_match.groups()
            lines = content.split('\n')
            question_text, options_list, is_question_part = "", [], True
            
            for line in lines:
                line = line.strip()
                if re.match(r'^[a-c]\)\s', line):
                    is_question_part = False
                    options_list.append(re.sub(r'^[a-c]\)\s*', '', line))
                elif is_question_part:
                    question_text += " " + line
                elif options_list:
                    options_list[-1] += " " + line
            
            if question_text and len(options_list) > 0:
                questions[q_num] = { "question": " ".join(question_text.strip().split()), "options": [opt.strip() for opt in options_list] }
        
        answer_pairs = re.findall(r'(\d+)\s+([A-C])', answer_content)
        answer_map = {num: letter for num, letter in answer_pairs}

        if not answer_map:
            logging.error("Could not parse the answer table.")
            return []

        final_data = []
        for q_num, q_data in questions.items():
            ans_letter = answer_map.get(q_num)
            if ans_letter:
                ans_index = ord(ans_letter.lower()) - 97
                if ans_index < len(q_data['options']):
                    q_data['answer'] = q_data['options'][ans_index]
                    final_data.append(q_data)
        
        logging.info(f"Parser 2 found {len(final_data)} questions.")
        return final_data

    except Exception as e:
        st.error(f"An error occurred during PDF parsing: {e}")
        logging.error(f"Parser 2 failed: {e}")
        return []

# ==============================================================================
# PARSER 3: ROBUST DOCX PARSER WITH BOLD DETECTION (Completely Rewritten)
# ==============================================================================
def parse_docx_with_bold_answers(docx_path):
    logging.info("LOG: Using NEW robust Parser 3 for DOCX.")
    try:
        doc = docx.Document(docx_path)
        questions_data = []
        current_q_block_paras = []

        for para in doc.paragraphs:
            para_text = para.text.strip()
            if not para_text or "technical test" in para_text.lower():
                continue

            # A new question starts with a number followed by a dot.
            if re.match(r'^\d{1,3}\.\s', para_text):
                if current_q_block_paras:
                    question = process_docx_block(current_q_block_paras)
                    if question:
                        questions_data.append(question)
                current_q_block_paras = [para]
            elif current_q_block_paras:
                current_q_block_paras.append(para)

        if current_q_block_paras: # Process the last block
            question = process_docx_block(current_q_block_paras)
            if question:
                questions_data.append(question)
        
        logging.info(f"Parser 3 found {len(questions_data)} questions.")
        return questions_data

    except Exception as e:
        st.error(f"Error reading DOCX file: {e}")
        logging.error(f"Parser 3 failed: {e}")
        return []

def process_docx_block(paragraphs):
    """
    Helper function to process a single block of DOCX paragraphs into one question.
    This version correctly handles multi-line questions and options.
    """
    if not paragraphs: return None

    question_text = ""
    options = []
    # This list will store dictionaries: {'text': "...", 'is_bold': False}
    structured_options = []

    is_question_part = True
    for para in paragraphs:
        line = para.text.strip()
        
        # Check if the line is the start of an option
        if re.match(r'^[a-zA-Z][.)]\s+', line):
            is_question_part = False # From now on, we are in the options part
            
            # Check if any text run in this paragraph is bold
            is_para_bold = any(run.bold for run in para.runs)
            clean_text = re.sub(r'^[a-zA-Z][.)]\s*', '', line).strip()
            structured_options.append({'text': clean_text, 'is_bold': is_para_bold})
            
        elif is_question_part:
            # If we haven't hit an option yet, this is part of the question
            question_text += " " + line
            
        elif structured_options:
            # If it's not a new option, it's a continuation of the previous option
            structured_options[-1]['text'] += " " + line
            # If any part of a multi-line option is bold, the whole option is the answer
            if any(run.bold for run in para.runs):
                structured_options[-1]['is_bold'] = True

    # Clean up the question text (remove the initial number)
    question_text = " ".join(re.sub(r'^\d{1,3}\.\s*', '', question_text).strip().split())
    
    # Extract the final data from our structured list
    options_list = [opt['text'] for opt in structured_options]
    correct_answer = next((opt['text'] for opt in structured_options if opt['is_bold']), None)

    if question_text and len(options_list) > 0 and correct_answer:
        return {
            "question": question_text,
            "options": options_list,
            "answer": correct_answer
        }
        
    logging.warning(f"Skipping invalid DOCX block. Question: '{question_text[:50]}...'")
    return None

# ==============================================================================
# STREAMLIT APPLICATION UI
# ==============================================================================
def main():
    st.set_page_config(page_title="Quiz Generator", page_icon="✈️", layout="centered")
    st.title("✈️ Universal Quiz Generator")

    if 'state' not in st.session_state:
        st.session_state.state = 'initial'
        st.session_state.questions = []
        st.session_state.current_question = 0
        st.session_state.user_answers = []

    if st.session_state.state == 'initial':
        st.info("Select your quiz format, then upload the corresponding file.")
        
        parser_choice = st.radio(
            "**1. Choose your file's format:**",
            ('**Pattern 1 (PDF):** Answer is `Ans. a` after each question.', 
             '**Pattern 2 (PDF):** Answers are in a table at the end.',
             '**Pattern 3 (DOCX):** The correct answer is **bolded**.')
        )

        file_type = "docx" if "Pattern 3" in parser_choice else "pdf"
        uploaded_file = st.file_uploader(f"**2. Upload your .{file_type} file:**", type=[file_type])

        if uploaded_file:
            temp_dir = "uploads"
            os.makedirs(temp_dir, exist_ok=True)
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            if 'Pattern 1' in parser_choice:
                st.session_state.questions = parse_pdf_with_ans_lines(temp_file_path)
            elif 'Pattern 2' in parser_choice:
                st.session_state.questions = parse_pdf_with_answer_table(temp_file_path)
            else: # Pattern 3
                st.session_state.questions = parse_docx_with_bold_answers(temp_file_path)
            
            if st.session_state.questions:
                st.session_state.state = 'quiz_started'
                st.rerun()
            else:
                st.error("Could not find any valid questions using the selected parser. Please check the file and your format selection.")
        return

    if st.session_state.state == 'finished':
        score = sum(1 for i, q in enumerate(st.session_state.questions) if i < len(st.session_state.user_answers) and st.session_state.user_answers[i] == q['answer'])
        
        st.success(f"## 🎯 Quiz Complete!")
        st.write(f"### Your Final Score: {score} out of {len(st.session_state.questions)}")
        with st.expander("Review Your Answers"):
             for i, q in enumerate(st.session_state.questions):
                user_answer = st.session_state.user_answers[i]
                correct_answer = q['answer']
                if user_answer == correct_answer:
                    st.markdown(f"**Q{i+1}: {q['question']}**\n\n✅ Your answer: `{user_answer}` (Correct)")
                else:
                    st.markdown(f"**Q{i+1}: {q['question']}**\n\n❌ Your answer: `{user_answer}`\n\nCorrect answer: `{correct_answer}`")
                st.markdown("---")
        if st.button("Take a New Quiz"):
            for key in list(st.session_state.keys()): del st.session_state[key]
            st.rerun()
        return

    if st.session_state.state in ['quiz_started', 'show_feedback']:
        q_index = st.session_state.current_question
        q_data = st.session_state.questions[q_index]

        st.subheader(f"Question {q_index + 1} of {len(st.session_state.questions)}")
        st.markdown(f"<p style='font-size: 20px; font-weight: 500;'>{q_data['question']}</p>", unsafe_allow_html=True)
        
        valid_options = [opt for opt in q_data["options"] if opt]
        
        if st.session_state.state == 'quiz_started':
            with st.form(key=f"form_{q_index}"):
                user_choice = st.radio("Choose your answer:", options=valid_options, key=f"radio_{q_index}")
                submitted = st.form_submit_button("Check Answer")
                if submitted:
                    st.session_state.user_answers.append(user_choice)
                    st.session_state.state = 'show_feedback'
                    st.rerun()
        
        elif st.session_state.state == 'show_feedback':
            last_answer = st.session_state.user_answers[q_index]
            try: default_index = valid_options.index(last_answer)
            except (ValueError, IndexError): default_index = 0
            
            st.radio("Your Answer:", options=valid_options, index=default_index, disabled=True)
            
            correct_answer = q_data['answer']
            if last_answer == correct_answer:
                st.success("✅ Correct!")
            else:
                st.error(f"❌ Incorrect! The correct answer was: **{correct_answer}**")
            
            button_text = "Finish Quiz" if q_index + 1 == len(st.session_state.questions) else "Next Question"
            if st.button(button_text):
                if q_index + 1 < len(st.session_state.questions):
                    st.session_state.current_question += 1
                    st.session_state.state = 'quiz_started'
                else:
                    st.session_state.state = 'finished'
                st.rerun()

if __name__ == "__main__":
    main()